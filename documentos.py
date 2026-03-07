import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

def get_pasta_padrao(pasta_escolhida=None):
    if pasta_escolhida and os.path.isdir(pasta_escolhida):
        return pasta_escolhida
    docs_path = Path.home() / "Documents" / "LexGen_Documentos"
    docs_path.mkdir(parents=True, exist_ok=True)
    return str(docs_path)

def gerar_peticao_docx(autor, reu, processo, acao, fatos, fundamentos, pedidos, valor, nome_escritorio, nome_advogado, oab, pasta_salvamento):
    doc = Document()
    
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(nome_escritorio.upper())
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold = True
    
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    
    doc.add_paragraph("EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA (...) VARA CÍVEL DA COMARCA DE (...).\n")
    
    # Adicionando o número do processo com destaque, caso exista
    if processo:
        p_proc = doc.add_paragraph(f"Processo nº: {processo}")
        p_proc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_proc.runs[0].bold = True
        doc.add_paragraph("\n")

    # Modifica a saudação se já tiver processo rolando
    saudacao = f"{autor.upper()}, já qualificado(a) nos autos em epígrafe, vem respeitosamente à presença de Vossa Excelência propor a presente:\n\n" if processo else f"{autor.upper()}, nacionalidade (...), estado civil (...), profissão (...), inscrito no CPF sob o nº (...), residente e domiciliado em (...), vem respeitosamente à presença de Vossa Excelência propor a presente:\n\n"
    corpo = doc.add_paragraph(saudacao)
    
    acao_p = doc.add_paragraph()
    acao_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    acao_run = acao_p.add_run(acao.upper() if acao else "AÇÃO JUDICIAL")
    acao_run.bold = True
    
    doc.add_paragraph(f"\nem face de {reu.upper()}, pelos fatos e fundamentos a seguir expostos:\n")

    doc.add_heading('I – DOS FATOS', level=1)
    doc.add_paragraph(fatos)
    doc.add_heading('II – DOS FUNDAMENTOS JURÍDICOS', level=1)
    doc.add_paragraph(fundamentos)
    doc.add_heading('III – DOS PEDIDOS', level=1)
    doc.add_paragraph("Diante do exposto, requer a Vossa Excelência:")
    doc.add_paragraph(pedidos)
    
    if valor:
        doc.add_heading('IV – DO VALOR DA CAUSA', level=1)
        doc.add_paragraph(f"Dá-se à causa o valor de {valor}.")

    # Assinatura dinâmica baseada no Cadastro
    adv_nome_formatado = nome_advogado.upper() if nome_advogado else "NOME DO ADVOGADO"
    doc.add_paragraph(f"\nNestes termos,\nPede Deferimento.\n\nLocal, Data.\n\n{adv_nome_formatado}\nADVOGADO(A)\nOAB {oab}")

    pasta_destino = get_pasta_padrao(pasta_salvamento)
    nome_base = f"{acao.replace(' ', '_')}_{autor.split()[0]}"
    caminho_docx = os.path.join(pasta_destino, f"{nome_base}.docx")
    caminho_pdf = os.path.join(pasta_destino, f"{nome_base}.pdf")
    
    doc.save(caminho_docx)
    try:
        convert(caminho_docx, caminho_pdf)
        return caminho_docx, caminho_pdf
    except Exception as e:
        return caminho_docx, None

def gerar_procuracao_docx(tipo, outorgante, outorgado, poderes, nome_escritorio, oab, pasta_salvamento):
    doc = Document()
    
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(nome_escritorio.upper())
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.bold = True
    
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run(f"PROCURAÇÃO {tipo.upper()}")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    
    doc.add_paragraph("\n")
    
    corpo = doc.add_paragraph()
    corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    corpo.add_run("OUTORGANTE: ").bold = True
    corpo.add_run(f"{outorgante}.\n\n")
    
    corpo.add_run("OUTORGADO(A): ").bold = True
    corpo.add_run(f"{outorgado}, advogado(a) inscrito(a) na OAB sob o nº {oab}, pertencente ao escritório {nome_escritorio}.\n\n")
    
    corpo.add_run("PODERES: ").bold = True
    corpo.add_run(f"{poderes}\n\n")
    
    doc.add_paragraph(f"Local e Data: _________________, ____ de ____________ de 20__.\n\n\n")
    
    assinatura = doc.add_paragraph("\n")
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nome_ass = doc.add_paragraph(outorgante)
    nome_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pasta_destino = get_pasta_padrao(pasta_salvamento)
    nome_base = f"Procuracao_{tipo.replace(' ', '_')}_{outorgante.split()[0]}"
    caminho_docx = os.path.join(pasta_destino, f"{nome_base}.docx")
    caminho_pdf = os.path.join(pasta_destino, f"{nome_base}.pdf")
    
    doc.save(caminho_docx)
    try:
        convert(caminho_docx, caminho_pdf)
        return caminho_docx, caminho_pdf
    except:
        return caminho_docx, None